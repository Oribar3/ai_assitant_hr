"""File handling utilities for resume processing."""

import os
from pathlib import Path
from typing import Optional, Tuple
from loguru import logger
import PyPDF2
import docx
from bs4 import BeautifulSoup


class FileHandler:
    """Handles file operations for resume processing."""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt', '.html'}
    
    @staticmethod
    def validate_file(file_path: str) -> Tuple[bool, str]:
        """
        Validate if file exists and has supported format.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        path = Path(file_path)
        
        if not path.exists():
            return False, f"File not found: {file_path}"
        
        if not path.is_file():
            return False, f"Path is not a file: {file_path}"
        
        if path.suffix.lower() not in FileHandler.SUPPORTED_EXTENSIONS:
            return False, f"Unsupported file format: {path.suffix}"
        
        return True, ""
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """
        Extract text from PDF file.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text content
        """
        try:
            text = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
            
            return "\n".join(text)
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {e}")
            raise
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text content
        """
        try:
            doc = docx.Document(file_path)
            text = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text.append(row_text)
            
            return "\n".join(text)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            raise
    
    @staticmethod
    def extract_text_from_txt(file_path: str) -> str:
        """
        Extract text from TXT file.
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            File content
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error reading TXT file {file_path}: {e}")
            raise
    
    @staticmethod
    def extract_text_from_html(file_path: str) -> str:
        """
        Extract text from HTML file.
        
        Args:
            file_path: Path to HTML file
            
        Returns:
            Extracted text content
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                soup = BeautifulSoup(file.read(), 'html.parser')
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                text = soup.get_text()
                # Clean up whitespace
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                return '\n'.join(chunk for chunk in chunks if chunk)
        except Exception as e:
            logger.error(f"Error extracting text from HTML {file_path}: {e}")
            raise
    
    @staticmethod
    def read_resume(file_path: str) -> str:
        """
        Read and extract text from resume file.
        
        Args:
            file_path: Path to resume file
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file doesn't exist
        """
        is_valid, error_msg = FileHandler.validate_file(file_path)
        if not is_valid:
            raise ValueError(error_msg)
        
        path = Path(file_path)
        extension = path.suffix.lower()
        
        logger.info(f"Reading resume from {file_path}")
        
        if extension == '.pdf':
            return FileHandler.extract_text_from_pdf(file_path)
        elif extension == '.docx':
            return FileHandler.extract_text_from_docx(file_path)
        elif extension == '.txt':
            return FileHandler.extract_text_from_txt(file_path)
        elif extension == '.html':
            return FileHandler.extract_text_from_html(file_path)
        else:
            raise ValueError(f"Unsupported file format: {extension}")
    
    @staticmethod
    def save_report(content: str, output_path: str) -> None:
        """
        Save report to file.
        
        Args:
            content: Report content
            output_path: Path to save the report
        """
        try:
            output_dir = Path(output_path).parent
            output_dir.mkdir(parents=True, exist_ok=True)
            
            with open(output_path, 'w', encoding='utf-8') as file:
                file.write(content)
            
            logger.info(f"Report saved to {output_path}")
        except Exception as e:
            logger.error(f"Error saving report to {output_path}: {e}")
            raise
